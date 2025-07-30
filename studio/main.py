import os
import uuid
from pathlib import Path
from typing import Dict, List

from fastapi import FastAPI, Request, Form, Depends, Query
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlmodel import Session
from starlette.middleware.sessions import SessionMiddleware
from docx import Document
import uvicorn

# Chat & model handling
from chat_client.generic_model_config import MODEL_REGISTRY, get_model_handler

# App components
from studio.routes import auth_routes, config_routes, chat_routes, newsdata_routes
from studio.services.db import init_db, get_session
from studio.dependency import get_current_user, require_roles
from studio.services.config_manager import load_user_config_db, save_user_config_db
# from studio.settings import DEFAULT_USER_CONFIG 

# Tool routes
from studio.auth import router as auth_router
from tools.context_manager.context_management_routes import router as context_router
from tools.RAG.rag_routes import router as rag_router
from tools.MARKDOWN.md_routes import router as md_router
from tools.DOCGEN.docgen_routes import router as docgen_router

app = FastAPI()

app.add_middleware(SessionMiddleware, secret_key="Shivaa@2025")
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
templates = Jinja2Templates(directory=str(Path(__file__).parent / "templates"))
app.mount("/static", StaticFiles(directory="studio/static"), name="static")
templates.env.cache.clear()


init_db()

# Include routers
app.include_router(auth_routes.router, tags=["Authentication"])     # Need special alignment.
app.include_router(config_routes.router, prefix="/config", tags=["Configuration"])
# app.include_router(google_router, prefix="/auth", tags=["OAuth"])
app.include_router(auth_router, prefix="/auth", tags=["OAuth"])     # Need special alignment.
app.include_router(newsdata_routes.router, tags=["News"])
app.include_router(chat_routes.router, prefix="/chat", tags=["Chat"])

# Configuration update
app.include_router(context_router, prefix="/tools/context_manager")
app.include_router(rag_router, prefix="/tools/rag")
app.include_router(md_router, prefix="/tools/markdown")
app.include_router(docgen_router, prefix="/tools/docgen")
    
conversations_store: Dict[str, List[Dict[str, str]]] = {}  # {conversation_name: [ {"user":..., "ai":...}, ... ]}

@app.post("/generate-doc")
async def generate_doc(request: Request, content: str = Form(...)):
    doc = Document()
    doc.add_heading("AI Response", level=1)
    doc.add_paragraph(content)

    filename = f"response_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(BASE_DIR, "generated", filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)

    return FileResponse(path=filepath, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {
        "request": request,
        "models": MODEL_REGISTRY
    })

@app.get("/me/config")
def get_config(
    user = Depends(get_current_user),
    session: Session = Depends(get_session),
):
    return load_user_config_db(session, user["id"], default={})

@app.put("/me/config")
def update_config(
    body: dict,
    user = Depends(get_current_user),
    session: Session = Depends(get_session),
):
    save_user_config_db(session, user["id"], body)
    return {"ok": True}

@app.get("/me")
def me(user = Depends(get_current_user)):
    return {"id": user.id, "email": user.email, "role": user.role}

@app.get("/admin/metrics")
def metrics(user = Depends(require_roles("admin"))):
    return {"secret": "..." }


@app.middleware("auth_context")
async def add_user_to_context(request: Request, call_next):
    user = None
    try:
        user = get_current_user(request)
    except:
        pass
    request.state.user = user
    response = await call_next(request)
    return response

@app.get("/chat/{model_id}", response_class=HTMLResponse)
async def get_chat_view(request: Request, model_id: str, conversation_name: str = Query(default=None)):
    conversation = conversations_store.get(conversation_name, []) if conversation_name else []

    return templates.TemplateResponse("chat.html", {
        "request": request,
        "agent": model_id,
        "response": None,  # No new response
        "user_input": None,
        "conversation_name": conversation_name,
        "conversation": conversation,
        "conversations": conversations_store,
        "models": MODEL_REGISTRY
    })

@app.post("/chat/{model_id}/send")
async def post_chat_message(
    model_id: str,
    user_input: str = Form(...),
    conversation_name: str = Form(...)
):
    handler = get_model_handler(model_id)
    ai_response = handler(user_input)

    # Store conversation
    if conversation_name not in conversations_store:
        conversations_store[conversation_name] = []
    conversations_store[conversation_name].append({
        "user": user_input,
        "ai": ai_response
    })

    return JSONResponse({
        "user": user_input,
        "ai": ai_response
    })


if __name__ == "__main__":
    uvicorn.run("studio.main:app", host="127.0.0.1", port=8030, reload=True)
